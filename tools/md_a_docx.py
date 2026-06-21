"""Conversor sencillo de Markdown a .docx (para generar la guía en Word).

Uso:  python tools\\md_a_docx.py GUIA.md GUIA.docx

Soporta: encabezados (#/##/###), listas con '-', listas numeradas, bloques de
código (```), citas/notas (>), y formato en línea **negrita** y `monoespaciado`.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


def _inline(parrafo, texto):
    for parte in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            parrafo.add_run(parte[2:-2]).bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            r = parrafo.add_run(parte[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            parrafo.add_run(parte)


def convertir(md_path: str, docx_path: str) -> str:
    lineas = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    en_codigo = False
    for linea in lineas:
        if linea.startswith("```"):
            en_codigo = not en_codigo
            continue
        if en_codigo:
            p = doc.add_paragraph()
            r = p.add_run(linea)
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_after = Pt(0)
            continue

        t = linea.rstrip()
        if not t:
            continue
        if t.startswith("# "):
            doc.add_heading(t[2:], level=0)
        elif t.startswith("## "):
            doc.add_heading(t[3:], level=1)
        elif t.startswith("### "):
            doc.add_heading(t[4:], level=2)
        elif t.strip() == "---":
            continue
        elif t.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, t.lstrip()[2:])
        elif re.match(r"^\s*\d+\.\s", t):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            _inline(p, t.strip())
        elif t.lstrip().startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            r_marker = p.add_run("")  # marcador vacío
            _inline(p, t.lstrip()[2:])
            for r in p.runs:
                r.italic = True
                if not r.font.color.rgb:
                    r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
        else:
            _inline(doc.add_paragraph(), t)

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python tools/md_a_docx.py <entrada.md> <salida.docx>")
        raise SystemExit(1)
    print("Generado:", convertir(sys.argv[1], sys.argv[2]))
